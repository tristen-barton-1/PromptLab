import os
import io
import time
import json
import html, csv
from io import BytesIO
from typing import List, Dict, Any

import pandas as pd
import streamlit as st

# Optional HTTP fallback (rare)
import httpx

# ================= Secrets & OpenAI bootstrap =================
if "OPENAI_API_KEY" in st.secrets and not os.getenv("OPENAI_API_KEY"):
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
API_KEY = os.getenv("OPENAI_API_KEY", "")

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False
    OpenAI = None

def ensure_client():
    if not OPENAI_AVAILABLE:
        raise RuntimeError("OpenAI SDK not installed in this environment.")
    return OpenAI(api_key=API_KEY if API_KEY else None)

# ================= Model options & helpers =================
MODEL_OPTIONS = [
    "gpt-5", "gpt-5-mini", "gpt-5-nano",
    "gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-3.5-turbo",
]
MODEL_ALIAS_MAP = {"ChatGPT 5": "gpt-5"}

def resolve_model_id(display_name: str) -> str:
    if not display_name:
        return "gpt-5"
    display_name = str(display_name).strip()
    return MODEL_ALIAS_MAP.get(display_name, display_name)

def is_gpt5(model_id: str) -> bool:
    return str(model_id or "").lower().startswith("gpt-5")

# ================= App chrome =================
st.set_page_config(page_title="Non-Technical Prompt Lab", page_icon="üß™", layout="wide")
st.title("üß™ Non-Technical Environment for Developing, Testing, and Refining Assistant Prompts")
st.caption("Write/refine **system instructions** and **prompts**, upload **one or more documents**, and test responses with **document context** ‚Äî no code required.")

# ================= Defaults: prompt rows (ONLY TWO) =================
DEFAULT_PROMPTS = pd.DataFrame([
    {
        "prompt": "How many documents can you see? List exact file names and a 1‚Äì2 line summary of each.",
        "system_instructions": "You are a document analysis assistant. Be concise and accurate in your responses. Always start with 'DOCUMENT COUNT:' followed by the number of documents.",
        "expected_result": "Should show 'DOCUMENT COUNT:' then all uploaded files with exact names & brief summaries.",
        "model": "", "backend": "", "temperature": "", "top_p": "", "max_tokens": ""
    },
    {
        "prompt": "Extract 3‚Äì5 key technical requirements per document. Quote exact lines with their section/heading if possible.",
        "system_instructions": "You are an RFP requirements specialist. Focus on technical specs and compliance requirements. Start with 'REQUIREMENTS FOUND:'.",
        "expected_result": "Quoted requirements per file, with exact filenames.",
        "model": "", "backend": "", "temperature": "", "top_p": "", "max_tokens": ""
    },
])

# ================= Session seeds =================
if "prompts_df" not in st.session_state:
    st.session_state["prompts_df"] = DEFAULT_PROMPTS.copy()
if "file_refs" not in st.session_state:
    st.session_state["file_refs"] = []  # [{'file_id','name','size'}]
if "results" not in st.session_state:
    st.session_state["results"] = []    # list of result dicts
if "export_format" not in st.session_state:
    st.session_state["export_format"] = ["CSV", "JSON"]

saved_df = st.session_state["prompts_df"]

def any_gpt5_in_df(df: pd.DataFrame) -> bool:
    try:
        for _, row in df.iterrows():
            m = str(row.get("model", "")).strip()
            if m and is_gpt5(resolve_model_id(m)):
                return True
    except Exception:
        pass
    return False

any_gpt5_from_saved = any_gpt5_in_df(saved_df)

# ================= Sidebar: Settings =================
with st.sidebar:
    st.header("Settings")

    if API_KEY:
        st.success("Using secret **OPENAI_API_KEY** (hidden).", icon="üîê")
    else:
        st.warning("Demo Mode ‚Äî no API key found in Secrets or env.", icon="‚ö†Ô∏è")

    backend = st.selectbox(
        "Backend",
        ["Completions", "Assistants", "Responses"],
        index=2,
        help=("‚Ä¢ **Completions**: Fast, single-turn tests (no retrieval/files).\n"
              "‚Ä¢ **Assistants**: Threads/Runs flow with message attachments.\n"
              "‚Ä¢ **Responses**: Unified API; takes input files directly.")
    )

    default_model_display = st.selectbox(
        "Default model (used when a row model is empty)",
        MODEL_OPTIONS,
        index=0
    )

    sampling_globally_disabled = is_gpt5(resolve_model_id(default_model_display)) or any_gpt5_from_saved

    st.markdown("**Generation defaults** (can be overridden per prompt):")
    if sampling_globally_disabled:
        st.info("Temperature & Top-p are hidden because a GPT-5 model is selected.", icon="üö´")
        temperature = 0.0
        top_p_default = 1.0
    else:
        temperature = st.slider("Temperature (default)", 0.0, 1.0, 0.2, 0.1)
        top_p_default = st.slider("Top-p (default)", 0.0, 1.0, 1.0, 0.05)

    max_output_tokens = st.slider("Max output tokens (default)", 64, 8192, 2048, 64)

    st.divider()
    st.write("**Export Options**")
    export_format = st.multiselect("Choose export format(s)", ["CSV", "JSON"], default=st.session_state["export_format"])
    st.session_state["export_format"] = export_format

# ================= Section 1: Instructions & Prompts =================
st.markdown("## 1) Instructions & Prompts")

system_instructions = st.text_area(
    "System Instructions",
    value=st.session_state.get(
        "system_instructions",
        "You are an expert RFP analysis assistant. Identify key requirements, deadlines, evaluation criteria, compliance obligations, and risks. Cite sections when possible and provide actionable recommendations. Always start your responses with 'GLOBAL ANALYSIS:' prefix.",
    ),
    height=120,
    key="system_instructions",
)

global_prompt = st.text_area(
    "Global Prompt",
    value=st.session_state.get(
        "global_prompt",
        "Provide a comprehensive analysis of the uploaded documents, focusing on key insights and actionable recommendations.",
    ),
    height=80,
    key="global_prompt",
)

# Prompt editor (stable schema; no column dropping)
REQUIRED_COLS = ["prompt", "system_instructions", "expected_result", "model", "backend", "max_tokens", "temperature", "top_p"]
for col in REQUIRED_COLS:
    if col not in st.session_state["prompts_df"].columns:
        st.session_state["prompts_df"][col] = ""

column_cfg = {
    "prompt": st.column_config.TextColumn("Prompt", width="medium"),
    "system_instructions": st.column_config.TextColumn("System Instructions (overrides global)", width="medium"),
    "expected_result": st.column_config.TextColumn("Expected result (notes)", width="medium"),
    "model": st.column_config.SelectboxColumn("Model (optional per prompt)", options=MODEL_OPTIONS),
    "backend": st.column_config.SelectboxColumn("Backend (optional per prompt)", options=["", "Completions", "Assistants", "Responses"]),
    "max_tokens": st.column_config.NumberColumn("Max tokens (optional)", min_value=64, max_value=8192, step=64),
}
if not sampling_globally_disabled:
    column_cfg["temperature"] = st.column_config.NumberColumn("Temp (0‚Äì1, optional)", min_value=0.0, max_value=1.0, step=0.1, format="%.1f")
    column_cfg["top_p"] = st.column_config.NumberColumn("Top-p (0‚Äì1, optional)", min_value=0.0, max_value=1.0, step=0.05, format="%.2f")

def _prepare_editor_df(df: pd.DataFrame, hide_sampling: bool) -> pd.DataFrame:
    df2 = df.copy()
    for col in ["temperature", "top_p", "max_tokens", "backend", "model", "system_instructions", "expected_result"]:
        if col not in df2.columns:
            df2[col] = ""
    if hide_sampling:
        return df2.drop(columns=[c for c in ["temperature", "top_p"] if c in df2.columns], errors="ignore")
    return df2

editor_input_df = _prepare_editor_df(st.session_state["prompts_df"], sampling_globally_disabled)
prompts_df_view = st.data_editor(
    editor_input_df,
    num_rows="dynamic",
    use_container_width=True,
    column_config=column_cfg,
    key="prompts_editor"
)

# Write back edits to the single source of truth
def _merge_back(edited: pd.DataFrame, original: pd.DataFrame) -> pd.DataFrame:
    # Simply use the edited dataframe as the new state
    # This allows proper row deletion and addition
    merged = edited.copy()
    
    # Ensure all required columns exist
    for c in REQUIRED_COLS:
        if c not in merged.columns:
            merged[c] = ""
    
    # Remove completely empty rows (all values are NaN or empty string)
    merged = merged.dropna(how='all')
    merged = merged[~(merged == '').all(axis=1)]
    
    return merged

st.session_state["prompts_df"] = _merge_back(prompts_df_view, st.session_state["prompts_df"])

# ================= Section 2: Upload files (no local parsing) =================
st.markdown("## 2) Upload Files (PDF, DOCX, TXT, etc.)")

uploaded = st.file_uploader(
    "",
    type=None,
    key="doc_uploader",
    accept_multiple_files=True
)

def upload_source_file(file_bytes: bytes, filename: str) -> str:
    """
    Upload a single file to OpenAI Files API, return file_id.
    Per OpenAI docs for file inputs: use purpose='user_data'.
    """
    client = ensure_client()
    f = client.files.create(file=(filename, io.BytesIO(file_bytes)), purpose="user_data")
    return f.id

if uploaded:
    files_list = uploaded if isinstance(uploaded, list) else [uploaded]
    for f in files_list:
        existing = next((r for r in st.session_state["file_refs"]
                         if r.get("name") == f.name and r.get("size") == getattr(f, "size", None)), None)
        if existing:
            continue
        f.seek(0)
        b = f.read()
        if API_KEY and OPENAI_AVAILABLE:
            try:
                with st.spinner(f"Uploading to OpenAI: {f.name}"):
                    fid = upload_source_file(b, f.name)
                st.session_state["file_refs"].append({"file_id": fid, "name": f.name, "size": getattr(f, "size", None)})
                st.success(f"Uploaded: {f.name}")
            except Exception as e:
                st.error(f"Upload failed for {f.name}: {e}")
        else:
            st.warning(f"Skipped upload (no API key): {f.name}")

file_refs = st.session_state.get("file_refs", [])


# ================= Helpers for model calls =================
def build_filename_preamble(files: List[Dict[str, Any]]) -> str:
    if not files:
        return ""
    names = "\n".join([f"- {f['name']}" for f in files])
    return (
        "You have the following files attached. Use and cite them by their exact filenames.\n"
        f"{names}\n\n"
        "When you quote or summarize, include the exact filename in parentheses like (filename.pdf).\n"
    )

def _gpt5_safe_payload(base: dict) -> dict:
    p = dict(base)
    p["reasoning"] = {"effort": "low"}
    p["text"] = {"verbosity": "low"}
    return p

def _responses_http_post(api_key: str, payload: dict) -> tuple[bool, dict | str]:
    url = "https://api.openai.com/v1/responses"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    try:
        with httpx.Client(timeout=120.0) as s:
            r = s.post(url, headers=headers, json=payload)
            if r.status_code >= 400:
                return False, r.text
            return True, r.json()
    except Exception as e:
        return False, str(e)

def _extract_responses_text_obj(resp) -> str:
    txt = getattr(resp, "output_text", None)
    if isinstance(txt, str) and txt.strip():
        return txt
    try:
        out = getattr(resp, "output", None)
        if out and len(out) > 0:
            content = getattr(out[0], "content", None)
            if content and len(content) > 0:
                item = content[0]
                t = getattr(item, "text", None)
                if hasattr(t, "value") and isinstance(t.value, str) and t.value.strip():
                    return t.value
                if isinstance(t, str) and t.strip():
                    return t
    except Exception:
        pass
    return ""

def _extract_responses_text_dict(data: dict) -> str | None:
    out_txt = data.get("output_text")
    if isinstance(out_txt, str) and out_txt.strip():
        return out_txt
    out = data.get("output")
    if isinstance(out, list) and out:
        content = out[0].get("content")
        if isinstance(content, list) and content:
            text_field = content[0].get("text")
            if isinstance(text_field, dict) and isinstance(text_field.get("value"), str):
                val = text_field.get("value", "")
                if val.strip():
                    return val
            if isinstance(text_field, str) and text_field.strip():
                return text_field
    return None

def call_openai_responses(system_instructions: str, prompt: str, model_id: str,
                          max_tokens: int, temp: float, top_p: float) -> str:
    """
    Responses API with file inputs using 'input_file' content items.
    All uploaded files are attached to every request.
    """
    files = st.session_state.get("file_refs") or []
    content: List[Dict[str, Any]] = [{"type": "input_file", "file_id": r["file_id"]} for r in files]

    preamble = build_filename_preamble(files)
    content.append({"type": "input_text", "text": f"{preamble}{prompt}"})

    payload: Dict[str, Any] = {
        "model": model_id,
        "instructions": system_instructions or "",
        "input": [{"role": "user", "content": content}],
        "max_output_tokens": int(max_tokens),
    }
    if is_gpt5(model_id):
        payload = _gpt5_safe_payload(payload)
    else:
        payload["temperature"] = float(temp)
        payload["top_p"] = float(top_p)

    if OPENAI_AVAILABLE and API_KEY:
        try:
            client = ensure_client()
            r = client.responses.create(**payload)
            text = _extract_responses_text_obj(r)
            if text and text.strip():
                return text
            if is_gpt5(model_id):
                payload_retry = dict(payload)
                payload_retry["max_output_tokens"] = min(int(max_tokens) + 1024, 8192)
                payload_retry["text"] = {"verbosity": "high"}
                r2 = client.responses.create(**payload_retry)
                text2 = _extract_responses_text_obj(r2)
                if text2 and text2.strip():
                    return text2
            return "Model returned no visible text. Try increasing **Max tokens**."
        except Exception:
            pass  # fall through to HTTP

    ok, data = _responses_http_post(API_KEY, payload)
    if ok and isinstance(data, dict):
        text = _extract_responses_text_dict(data)
        if text and str(text).strip():
            return text
        return "Model returned no visible text. Try increasing **Max tokens**."
    elif not ok:
        return f"HTTP request failed: {data}"
    return "Model returned no visible text."

def call_openai_completions(messages, model_id: str, max_tokens: int, temp: float, top_p: float):
    # No retrieval/files in chat completions. Route GPT-5 to Responses.
    if is_gpt5(model_id):
        return ("‚ö†Ô∏è Auto-routed: GPT-5 is not available through Chat Completions; using Responses instead.\n\n" +
                call_openai_responses(messages[0]["content"], messages[-1]["content"], model_id, max_tokens, temp, top_p))
    if not OPENAI_AVAILABLE:
        raise RuntimeError("OpenAI SDK not installed in this environment.")
    client = ensure_client()
    resp = client.chat.completions.create(
        model=model_id, messages=messages, max_tokens=max_tokens, temperature=temp, top_p=top_p,
    )
    return resp.choices[0].message.content

def call_openai_assistants(system_instructions: str, prompt: str, model_id: str) -> str:
    """
    Assistants v2 with message-level attachments. All uploaded files are attached to the user message.
    """
    if not OPENAI_AVAILABLE:
        raise RuntimeError("OpenAI SDK not installed in this environment.")
    client = ensure_client()

    assistant = client.beta.assistants.create(
        name="BidWERX Prompt Lab",
        instructions=system_instructions or "",
        model=model_id,
        tools=[{"type": "file_search"}],
    )

    thread = client.beta.threads.create()

    files = st.session_state.get("file_refs") or []
    msg_attachments = [{"file_id": r["file_id"], "tools": [{"type": "file_search"}]} for r in files]

    preamble = build_filename_preamble(files)
    user_content = [{"type": "input_text", "text": f"{preamble}{prompt}"}]

    client.beta.threads.messages.create(
        thread_id=thread.id,
        role="user",
        content=user_content,
        attachments=msg_attachments if msg_attachments else None,
    )

    run = client.beta.threads.runs.create(thread_id=thread.id, assistant_id=assistant.id)
    for _ in range(60):
        run = client.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
        if run.status == "completed":
            break
        if run.status in ("failed", "cancelled", "expired"):
            raise RuntimeError(f"Assistant run ended with status: {run.status}")
        time.sleep(1)

    msgs = client.beta.threads.messages.list(thread_id=thread.id, order="desc", limit=1)
    if msgs.data:
        out = []
        for p in msgs.data[0].content:
            if getattr(p, "type", None) == "text":
                out.append(p.text.value)
        if out:
            return "\n".join(out)
    return "(No response returned from Assistants API)"

# ================= Utilities =================
def _coalesce_num(val, default, cast=float):
    try:
        if pd.isna(val) or str(val).strip() == "":
            return cast(default)
        return cast(val)
    except Exception:
        return cast(default)

def _safe_str(val, default=""):
    """Safely convert a value to string, handling NaN and None values."""
    if pd.isna(val) or val is None:
        return default
    return str(val).strip()

def render_card(container: st.delta_generator.DeltaGenerator, result: Dict[str, Any], expand: bool = False):
    """Stable card renderer (single format)."""
    files_line = result.get("documents") or "‚Äî"
    title = (result.get("prompt") or "(no prompt)")[:120]
    with container.expander(f"üß™ {title}", expanded=expand):
        st.markdown(f"**Files:** {files_line}")
        st.markdown(
            f"**Backend:** {result.get('backend','')}  |  **Model:** {result.get('model_label') or result.get('model_id','')}  \n"
            f"**Temp:** {result.get('temperature_used','')}  |  **Top-p:** {result.get('top_p_used','')}  |  "
            f"**Max tokens:** {result.get('max_tokens_used','')}"
        )
        st.markdown("**AI Response:**")
        st.markdown(result.get("ai_response",""))

# ================= Section 3: Run Tests & Results (single live/persistent area) =================
st.markdown("## 3) Run Tests & Results")

# This is the ONLY place results are rendered (both during runs and on reruns)
cards_root = st.container()

# Show button only when not processing
run_clicked = False
button_placeholder = st.empty()

if not st.session_state.get("processing", False):
    run_clicked = button_placeholder.button("‚ñ∂Ô∏è Run Prompts", type="primary", use_container_width=True)
    if run_clicked:
        st.session_state["processing"] = True
        button_placeholder.empty()  # Clear the button immediately
else:
    button_placeholder.empty()  # Clear the button when processing

if run_clicked or st.session_state.get("processing", False):
    # Reset results only when starting a new run
    st.session_state["results"] = []

    # Validate
    files = st.session_state.get("file_refs", [])
    if not system_instructions.strip():
        st.warning("Please provide system instructions.")
        st.session_state["processing"] = False
    elif st.session_state["prompts_df"].empty and not global_prompt.strip():
        st.warning("Please provide at least one prompt in the table or a global prompt.")
        st.session_state["processing"] = False
    elif not files:
        st.warning("Please upload at least one file for grounding.")
        st.session_state["processing"] = False
    else:
        # Freeze work items from the edited table
        df_now = st.session_state["prompts_df"]
        work_items: List[Dict[str, Any]] = []
        for _, row in df_now.iterrows():
            prompt_text = _safe_str(row.get("prompt")) or global_prompt.strip()
            if not prompt_text:
                continue
            work_items.append({
                "prompt": prompt_text,
                "expected_result": _safe_str(row.get("expected_result")),
                "model_display": _safe_str(row.get("model")) or default_model_display,
                "backend": _safe_str(row.get("backend")) or backend,
                "system_instructions": _safe_str(row.get("system_instructions")) or system_instructions,
                "temperature": _coalesce_num(row.get("temperature"), 0.2 if not sampling_globally_disabled else 0.0, float),
                "top_p": _coalesce_num(row.get("top_p"), 1.0, float),
                "max_tokens": _coalesce_num(row.get("max_tokens"), max_output_tokens, int),
            })

        total = len(work_items)
        if total == 0:
            st.warning("No valid prompts found to process.")
        else:
            # Progress indicator visible immediately
            progress_ph = st.empty()
            progress_ph.progress(0, text=f"Started‚Ä¶ (0 of {total})")

            # Pre-create placeholders for live cards inside the SINGLE cards area
            card_placeholders = [cards_root.container() for _ in range(total)]

            # Process prompts sequentially and render each card as it completes
            for i, item in enumerate(work_items, start=1):
                time.sleep(0.03)  # small yield for UI

                model_id = resolve_model_id(item["model_display"])
                used_backend = item["backend"]
                notes = []
                if used_backend == "Completions" and is_gpt5(model_id):
                    used_backend = "Responses"
                    notes.append("‚ö†Ô∏è Auto-routed from Completions to Responses for GPT-5.")
                if is_gpt5(model_id):
                    notes.append("‚ö†Ô∏è GPT-5: Temperature / Top-p are not applied.")

                try:
                    if used_backend == "Assistants":
                        ai_text = call_openai_assistants(item["system_instructions"], item["prompt"], model_id)
                    elif used_backend == "Responses":
                        ai_text = call_openai_responses(
                            item["system_instructions"], item["prompt"], model_id,
                            item["max_tokens"], item["temperature"], item["top_p"]
                        )
                    else:
                        messages = [
                            {"role": "system", "content": item["system_instructions"] or ""},
                            {"role": "user", "content": item["prompt"]},
                        ]
                        ai_text = call_openai_completions(
                            messages, model_id, item["max_tokens"], item["temperature"], item["top_p"]
                        )
                    if notes:
                        ai_text = ("\n\n".join(notes) + "\n\n" + (ai_text or "")).strip()
                except Exception as e:
                    ai_text = f"ERROR: {e}"

                # Persist and immediately render this card
                result_row = {
                    "documents": ", ".join([r["name"] for r in files]) if files else "",
                    "system_instructions": item["system_instructions"],
                    "prompt": item["prompt"],
                    "expected_result": item["expected_result"],
                    "backend": used_backend,
                    "model_label": item["model_display"],
                    "model_id": model_id,
                    "temperature_used": "(ignored for GPT-5)" if is_gpt5(model_id) else item["temperature"],
                    "top_p_used": "(ignored for GPT-5)" if is_gpt5(model_id) else item["top_p"],
                    "max_tokens_used": item["max_tokens"],
                    "ai_response": ai_text
                }
                st.session_state["results"].append(result_row)

                render_card(card_placeholders[i-1], result_row, expand=(i == 1))

                pct = int((i / total) * 100)
                progress_ph.progress(pct, text=f"Processed {i} of {total}")

            progress_ph.progress(100, text=f"Completed {total} of {total}")
            # Clear processing flag when done
            st.session_state["processing"] = False

else:
    # No new run: render whatever is already in session (single area, no duplicates)
    existing_results = st.session_state.get("results", [])
    if existing_results:
        for res in existing_results:
            slot = cards_root.container()
            render_card(slot, res, expand=False)

# ======= Downloads =======
files_for_view = st.session_state.get("file_refs", [])
results = st.session_state.get("results", [])
if results:
    st.markdown("### Download Results")

    def build_markdown_report(title, system_instructions, files, results):
        lines = [
            f"# {title}", "",
            f"**Files:** {', '.join([f['name'] for f in files]) if files else '‚Äî'}", "",
            "## System Instructions", "",
            system_instructions or "_None_", "",
            "## Results", "",
        ]
        for i, r in enumerate(results, 1):
            lines += [
                f"### {i}. {r.get('prompt','')}", "",
                (f"*Backend:* {r.get('backend','')} | *Model:* {r.get('model_label','') or r.get('model_id','')} | "
                 f"*Temp:* {r.get('temperature_used','')} | *Top-p:* {r.get('top_p_used','')} | "
                 f"*Max tokens:* {r.get('max_tokens_used','')}"),
                "",
                "**AI Response**", "",
                r.get("ai_response",""), "",
            ]
            if r.get("expected_result"):
                lines += ["**Expected Result (notes)**", "", r["expected_result"], ""]
        return "\n".join(lines)

    def build_html_report(title, system_instructions, files, results):
        css = """
        <style>
          body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial;line-height:1.55;color:#222;
               max-width:980px;margin:2rem auto;padding:0 1rem;}
          h1,h2,h3{margin:.4rem 0 .6rem}
          .meta{color:#555;margin:.3rem 0 1rem;font-size:.95rem}
          .card{border:1px solid #eaecef;border-radius:12px;padding:16px;margin:16px 0;
                box-shadow:0 1px 2px rgba(0,0,0,.04)}
          pre{white-space:pre-wrap;background:#f6f8fa;border:1px solid #eaecef;border-radius:8px;
              padding:12px;margin:.5rem 0}
          .small{font-size:.9rem;color:#444}
        </style>
        """
        def esc(x):
            import html as _h
            return _h.escape(x or "")
        file_names = ', '.join([f['name'] for f in files]) if files else '‚Äî'
        parts = [f"<h1>{esc(title)}</h1>",
                 f"<div class='meta'><b>Files:</b> {esc(file_names)}</div>",
                 "<h2>System Instructions</h2>",
                 f"<pre>"+esc(system_instructions or 'None')+"</pre>",
                 "<h2>Results</h2>"]
        for i, r in enumerate(results, 1):
            parts.append("<div class='card'>")
            parts.append(f"<h3>{i}. {esc(r.get('prompt',''))}</h3>")
            parts.append(
                f"<div class='small'><b>Backend:</b> {esc(r.get('backend',''))} &nbsp; "
                f"<b>Model:</b> {esc(r.get('model_label','') or r.get('model_id',''))} &nbsp; "
                f"<b>Temp:</b> {esc(str(r.get('temperature_used','')))} &nbsp; "
                f"<b>Top-p:</b> {esc(str(r.get('top_p_used','')))} &nbsp; "
                f"<b>Max tokens:</b> {esc(str(r.get('max_tokens_used','')))}</div>"
            )
            exp = r.get("expected_result") or ""
            if exp:
                parts.append(f"<div class='small'><b>Expected Result (notes):</b> {esc(exp)}</div>")
            parts.append("<h4>AI Response</h4>")
            parts.append("<pre>"+esc(r.get('ai_response',''))+"</pre>")
            parts.append("</div>")
        html_doc = f"<!doctype html><html><head><meta charset='utf-8'><title>{esc(title)}</title>{css}</head><body>{''.join(parts)}</body></html>"
        return html_doc.encode("utf-8")

    def build_docx_report(title, system_instructions, files, results):
        from io import BytesIO as _BytesIO
        from docx import Document
        d = Document()
        d.add_heading(title, 0)
        file_names = ', '.join([f['name'] for f in files]) if files else '‚Äî'
        d.add_paragraph(f"Files: {file_names}")
        d.add_heading("System Instructions", level=1)
        d.add_paragraph(system_instructions or "None")
        d.add_heading("Results", level=1)
        for i, r in enumerate(results, 1):
            d.add_heading(f"{i}. {r.get('prompt','')}", level=2)
            d.add_paragraph(
                f"Backend: {r.get('backend','')} | Model: {r.get('model_label','') or r.get('model_id','')} | "
                f"Temp: {r.get('temperature_used','')} | Top-p: {r.get('top_p_used','')} | "
                f"Max tokens: {r.get('max_tokens_used','')}"
            )
            d.add_paragraph("AI Response:")
            d.add_paragraph(r.get("ai_response",""))
            if r.get("expected_result"):
                d.add_paragraph("Expected Result (notes):")
                d.add_paragraph(r["expected_result"])
        bio = _BytesIO()
        d.save(bio)
        return bio.getvalue()

    md_text   = build_markdown_report("BidWERX Prompt Lab ‚Äî Test Run", st.session_state.get("system_instructions"), files_for_view, results)
    html_bytes = build_html_report("BidWERX Prompt Lab ‚Äî Test Run", st.session_state.get("system_instructions"), files_for_view, results)
    docx_bytes = build_docx_report("BidWERX Prompt Lab ‚Äî Test Run", st.session_state.get("system_instructions"), files_for_view, results)

    export_format = st.session_state.get("export_format", ["CSV", "JSON"])
    if "CSV" in export_format:
        csv_bytes = pd.DataFrame(results).to_csv(index=False).encode("utf-8")
        st.download_button("‚¨áÔ∏è Download CSV", data=csv_bytes, file_name="prompt_lab_results.csv", mime="text/csv")
    if "JSON" in export_format:
        json_bytes = pd.DataFrame(results).to_json(orient="records", indent=2).encode("utf-8")
        st.download_button("‚¨áÔ∏è Download JSON", data=json_bytes, file_name="prompt_lab_results.json", mime="application/json")

    st.download_button("‚¨áÔ∏è Download HTML Report", data=html_bytes, file_name="prompt_lab_report.html", mime="text/html")
    st.download_button("‚¨áÔ∏è Download Markdown", data=md_text.encode("utf-8"), file_name="prompt_lab_report.md", mime="text/markdown")
    st.download_button("‚¨áÔ∏è Download DOCX (Word)", data=docx_bytes,
                       file_name="prompt_lab_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
